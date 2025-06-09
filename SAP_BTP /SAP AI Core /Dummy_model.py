# attachment_analyzer.py

import os
import logging
import base64
import zipfile
import tempfile
import shutil
import fitz  # PyMuPDF for PDF processing
import docx
import pandas as pd
from openpyxl import load_workbook
from langchain_core.messages import HumanMessage
from src.models.llm_models import LLM_Models
from typing import List, Dict, Any


def create_support_analysis_prompt(summary: str, description: str) -> str:
    """
    Build the prompt that asks the LLM to analyze an attached image/file
    for relevance against the given ticket summary & description.
    """
    prompt = (
        "You are a technical support agent analyzing an attached image to determine its relevance to the reported issue.\n\n"
        f"**Issue Summary:** {summary}\n\n"
        f"**Issue Description:** {description}\n\n"
        "**Your Task:**\n"
        "1. Analyze the attached image in detail\n"
        "2. Compare the image content with the issue summary and description\n"
        "3. Determine if the attachment is relevant to resolving this specific issue\n\n"
        "**Analysis Guidelines:**\n"
        "- **For Technical Images (Screenshots, Error Messages, System Interfaces):**\n"
        "  • Describe visible UI elements, error dialogs, system screens, or application interfaces\n"
        "  • Identify application/module names, error codes, warning messages\n"
        "  • Note any technical details like timestamps, file paths, user IDs, or system states\n"
        "  • Look for access-related indicators (login prompts, permission errors, authentication issues)\n"
        "  • Mention any visible user actions or workflow steps\n\n"
        "- **For Non-Technical Images:**\n"
        "  • Provide a brief description of the image content\n"
        "  • Explain why it may or may not relate to the reported issue\n\n"
        "**Relevance Assessment:**\n"
        "After analyzing the image, determine similarity and relevance by considering:\n"
        "- Does the image show the same system/application mentioned in the issue?\n"
        "- Are there matching error messages, codes, or symptoms?\n"
        "- Does the image demonstrate the problem described?\n"
        "- Would this image help a support engineer understand or resolve the issue?\n"
        "- Are there related workflow steps, data, or system components visible?\n\n"
        "**Response Format:**\n"
        "1. **Image Description:** [Detailed description of what the image shows]\n"
        "2. **Technical Details:** [Any error codes, system names, technical elements - if none, state 'No technical elements detected']\n"
        "3. **Relevance Analysis:** [Explain the connection or lack thereof between image and issue]\n"
        "4. **Final Assessment:** **RELEVANT** or **NOT RELEVANT**\n"
        "5. **Support Reasoning:** [Brief explanation of why this assessment helps in issue resolution]\n\n"
        "End your response with the word: `TERMINATE`"
    )
    return prompt


class AttachmentAnalyzer:
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}
    DOCUMENT_EXTENSIONS = {'.pdf', '.doc', '.docx'}
    SPREADSHEET_EXTENSIONS = {'.xls', '.xlsx', '.csv'}
    ZIP_EXTENSION = '.zip'

    @staticmethod
    def encode_file_to_base64(file_path: str) -> str:
        try:
            with open(file_path, 'rb') as f:
                return base64.b64encode(f.read()).decode('utf-8')
        except Exception as e:
            logging.error(f"Error encoding file: {e}")
            return None

    @staticmethod
    def decode_and_save_temp(base64_str: str, filename: str) -> str:
        path = f"/tmp/{filename}"
        with open(path, "wb") as f:
            f.write(base64.b64decode(base64_str))
        return path

    @staticmethod
    def _get_file_mime_type(file_path: str) -> str:
        """Determine the MIME type based on file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in AttachmentAnalyzer.IMAGE_EXTENSIONS:
            return "image/png" if ext == ".png" else "image/jpeg"
        if ext == '.pdf':
            return "application/pdf"
        if ext in {'.doc', '.docx'}:
            return "application/msword"
        if ext in {'.xls', '.xlsx'}:
            return "application/vnd.ms-excel"
        if ext == '.csv':
            return "text/csv"
        if ext == AttachmentAnalyzer.ZIP_EXTENSION:
            return "application/zip"
        return "application/octet-stream"

    @staticmethod
    def process_attachment(
        attachment: str,
        summary: str,
        description: str
    ) -> (bool, Any):
        """
        Process one attachment and return (is_relevant: bool, summary_or_structured).
        Pass in summary and description separately.
        """
        # 1) Direct multimodal attempt
        try:
            valid, response_text = AttachmentAnalyzer._try_direct_multimodal(
                attachment, summary, description
            )
            if valid:
                return True, response_text
        except Exception as e:
            logging.warning(f"Direct LLM approach failed for {attachment}: {e}")

        # 2) Fallback based on extension
        ext = os.path.splitext(attachment)[1].lower()
        if ext in AttachmentAnalyzer.IMAGE_EXTENSIONS:
            return AttachmentAnalyzer._process_simple_image(
                attachment, summary, description
            )
        elif ext == '.pdf':
            return AttachmentAnalyzer._process_pdf_with_pymupdf(
                attachment, summary, description
            )
        elif ext in {'.doc', '.docx'}:
            return AttachmentAnalyzer._process_word_document(
                attachment, summary, description
            )
        elif ext in {'.xls', '.xlsx'}:
            return AttachmentAnalyzer._process_excel_document(
                attachment, summary, description
            )
        elif ext == '.csv':
            return AttachmentAnalyzer._process_with_text_extraction(
                attachment, summary, description
            )
        elif ext == AttachmentAnalyzer.ZIP_EXTENSION:
            return AttachmentAnalyzer._process_zip_attachment(
                attachment, summary, description
            )
        else:
            return False, f"Unsupported file type: {ext}"

    @staticmethod
    def _process_simple_image(
        image_path: str,
        summary: str,
        description: str
    ) -> (bool, str):
        """
        1) Build the support-analysis prompt using summary & description directly.
        2) Send prompt + base64 image to LLM.
        3) Parse "Final Assessment: RELEVANT" vs. "NOT RELEVANT".
        """
        # 1) Build the LLM prompt using summary & description
        prompt = create_support_analysis_prompt(summary, description)

        # 2) Encode image as base64 + get MIME type
        b64 = AttachmentAnalyzer.encode_file_to_base64(image_path)
        if not b64:
            return False, "Error encoding image file."

        mime = AttachmentAnalyzer._get_file_mime_type(image_path)
        if not mime.startswith("image/"):
            return False, f"Unsupported image type: {mime}"

        # 3) Invoke LLM with both prompt text and image payload
        llm, _ = LLM_Models.initialize_models()
        msg = HumanMessage(content=[
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
        ])
        try:
            res = llm.invoke([msg])
        except Exception as e:
            logging.error(f"LLM call failed: {e}")
            return False, f"LLM call error: {e}"

        full_response = getattr(res, 'content', str(res)).strip()

        # 4) Parse "Final Assessment: RELEVANT" vs. "NOT RELEVANT"
        upper = full_response.upper()
        if "FINAL ASSESSMENT: RELEVANT" in upper:
            return True, full_response
        elif "FINAL ASSESSMENT: NOT RELEVANT" in upper:
            return False, full_response
        else:
            # Model didn't follow instructions exactly → default to "not relevant"
            return False, full_response

    @staticmethod
    def _is_relevant_text(attachment_summary: str, ticket_text: str) -> bool:
        """
        (No longer used when relying fully on prompt-based relevance,
         but left here if you want a fallback logic.)
        Compare the generated attachment_summary vs. ticket_text.
        """
        from difflib import SequenceMatcher

        a = (attachment_summary or "").lower().strip()
        b = (ticket_text or "").lower().strip()
        if not a or not b:
            return False

        ratio = SequenceMatcher(None, a, b).ratio()
        return ratio >= 0.60

    @staticmethod
    def _process_zip_attachment(
        zip_path: str,
        summary: str,
        description: str
    ) -> (bool, List[Dict[str, Any]]):
        """
        Extract a zip file and process each contained file.
        Returns:
          - overall_valid: True if any file is relevant
          - structured_summaries: List of dicts with filename, relevant flag, and summary
        """
        try:
            temp_dir = tempfile.mkdtemp()
            with zipfile.ZipFile(zip_path, 'r') as z:
                z.extractall(temp_dir)

            overall_valid = False
            structured_summaries: List[Dict[str, Any]] = []

            for root, _, files in os.walk(temp_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    valid, summary_or_string = AttachmentAnalyzer.process_attachment(
                        full_path, summary, description
                    )

                    structured_summaries.append({
                        "filename": file,
                        "relevant": valid,
                        "summary": summary_or_string
                    })

                    if valid:
                        overall_valid = True

            return overall_valid, structured_summaries

        except Exception as e:
            logging.error(f"Error processing zip file: {e}")
            return False, [{"filename": os.path.basename(zip_path),
                            "relevant": False,
                            "summary": f"Error processing zip: {e}"}]

    @staticmethod
    def _process_with_text_extraction(
        file_path: str,
        summary: str,
        description: str
    ) -> (bool, str):
        """Process a document by extracting text and analyzing it."""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            text = ""
            if ext in {'.doc', '.docx'}:
                text = "\n".join([p.text for p in docx.Document(file_path).paragraphs])
            elif ext in {'.xls', '.xlsx'}:
                text = pd.read_excel(file_path).to_string(index=False)
            elif ext == '.csv':
                text = pd.read_csv(file_path).to_string(index=False)
            else:
                return False, f"Text extraction not supported for {ext}"

            if not text.strip():
                return False, f"No text extracted from {ext}"

            # Build a prompt similar to the image prompt, but for document text
            prompt = (
                "You are a technical support agent analyzing an attached document to determine its relevance to the reported issue.\n\n"
                f"**Issue Summary:** {summary}\n\n"
                f"**Issue Description:** {description}\n\n"
                "**Your Task:**\n"
                "1. Read the extracted document text in detail\n"
                "2. Compare the document content with the issue summary and description\n"
                "3. Determine if the document is relevant to resolving this specific issue\n\n"
                "**Response Format:**\n"
                "1. **Document Summary:** [Brief summary of relevant parts, or 'No relevant content found']\n"
                "2. **Final Assessment:** **RELEVANT** or **NOT RELEVANT**\n"
                "3. **Support Reasoning:** [Why this decision helps in issue resolution]\n\n"
                "End your response with the word: `TERMINATE`"
            )

            llm, _ = LLM_Models.initialize_models()
            msg = HumanMessage(content=prompt + "\n\n" + text[:3000])
            res = llm.invoke([msg])
            analysis = getattr(res, 'content', str(res)).strip()

            upper = analysis.upper()
            if "FINAL ASSESSMENT: RELEVANT" in upper:
                return True, analysis
            elif "FINAL ASSESSMENT: NOT RELEVANT" in upper:
                return False, analysis
            else:
                return False, analysis

        except Exception as e:
            logging.error(f"Error extracting text: {e}")
            return False, f"Error extracting text: {e}"

    @staticmethod
    def _process_pdf_with_pymupdf(
        pdf_path: str,
        summary: str,
        description: str
    ) -> (bool, str):
        """
        Process PDF with PyMuPDF to extract both text and images,
        then analyze with multi-modal LLM + relevance prompt.
        """
        temp_dir = tempfile.mkdtemp()
        imgs = []
        try:
            doc = fitz.open(pdf_path)
            text = "".join([page.get_text() + "\n" for page in doc])

            # Extract embedded images
            for page in doc:
                for img in page.get_images(full=True):
                    xref = img[0]
                    base = doc.extract_image(xref)
                    pth = os.path.join(temp_dir, f"img_{xref}.{base['ext']}")
                    with open(pth, 'wb') as f:
                        f.write(base['image'])
                    imgs.append(pth)

            # 1) Summarize extracted text + assess relevance
            prompt_text = (
                "You are a technical support agent analyzing an attached PDF to determine its relevance to the reported issue.\n\n"
                f"**Issue Summary:** {summary}\n\n"
                f"**Issue Description:** {description}\n\n"
                "**Your Task:**\n"
                "1. Read the extracted PDF text in detail\n"
                "2. Compare the PDF content with the issue summary and description\n"
                "3. Determine if the PDF (text + any images) is relevant\n\n"
                "**Response Format:**\n"
                "1. **PDF Text Summary:** [Brief summary of relevant parts, or 'No relevant content found']\n"
                "2. **Final Assessment:** **RELEVANT** or **NOT RELEVANT**\n"
                "3. **Support Reasoning:** [Why this decision helps in issue resolution]\n\n"
                "End your response with the word: `TERMINATE`"
            )

            combined_text_result = ""
            if text.strip():
                llm, _ = LLM_Models.initialize_models()
                msg_text = HumanMessage(content=prompt_text + "\n\n" + text[:3000])
                res_text = llm.invoke([msg_text])
                combined_text_result = getattr(res_text, 'content', '').strip()

                upper_text = combined_text_result.upper()
                if "FINAL ASSESSMENT: RELEVANT" in upper_text:
                    return True, combined_text_result
                elif "FINAL ASSESSMENT: NOT RELEVANT" in upper_text:
                    return False, combined_text_result

            # 2) If text didn’t decide, analyze each embedded image
            for img_path in imgs:
                is_valid, img_summary = AttachmentAnalyzer._process_simple_image(
                    img_path, summary, description
                )
                if is_valid:
                    return True, img_summary

            # 3) If neither text nor any images were flagged “relevant”, return the text result or a fallback
            if combined_text_result:
                return False, combined_text_result
            else:
                return False, "No relevant content found in PDF."

        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    @staticmethod
    def _process_word_document(
        doc_path: str,
        summary: str,
        description: str
    ) -> (bool, str):
        """
        Process Word documents including extraction of embedded images and ticket-specific analysis.
        """
        temp_dir = tempfile.mkdtemp()
        imgs = []
        try:
            # Extract text from paragraphs and tables
            doc = docx.Document(doc_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            # Extract embedded images
            pkg = docx.Document(doc_path).part.package
            idx = 0
            for rel in pkg.main_document_part.rels.values():
                if rel.target_part and rel.target_part.content_type.startswith('image/'):
                    ext = rel.target_part.content_type.split('/')[1]
                    pth = os.path.join(temp_dir, f"img_{idx}.{ext}")
                    with open(pth, 'wb') as f:
                        f.write(rel.target_part.blob)
                    imgs.append(pth)
                    idx += 1

            # 1) Summarize extracted text + assess relevance
            prompt_text = (
                "You are a technical support agent analyzing an attached Word document to determine its relevance to the reported issue.\n\n"
                f"**Issue Summary:** {summary}\n\n"
                f"**Issue Description:** {description}\n\n"
                "**Your Task:**\n"
                "1. Read the extracted document text in detail\n"
                "2. Compare the document content with the issue summary and description\n"
                "3. Determine if the document (text + any embedded images) is relevant\n\n"
                "**Response Format:**\n"
                "1. **Document Text Summary:** [Brief summary of relevant parts, or 'No relevant content found']\n"
                "2. **Final Assessment:** **RELEVANT** or **NOT RELEVANT**\n"
                "3. **Support Reasoning:** [Why this decision helps in issue resolution]\n\n"
                "End your response with the word: `TERMINATE`"
            )

            combined_text_result = ""
            if text.strip():
                llm, _ = LLM_Models.initialize_models()
                msg_text = HumanMessage(content=prompt_text + "\n\n" + text[:3000])
                res_text = llm.invoke([msg_text])
                combined_text_result = getattr(res_text, 'content', '').strip()

                upper_text = combined_text_result.upper()
                if "FINAL ASSESSMENT: RELEVANT" in upper_text:
                    return True, combined_text_result
                elif "FINAL ASSESSMENT: NOT RELEVANT" in upper_text:
                    return False, combined_text_result

            # 2) If text didn’t decide, analyze each embedded image
            for img_path in imgs:
                is_valid, img_summary = AttachmentAnalyzer._process_simple_image(
                    img_path, summary, description
                )
                if is_valid:
                    return True, img_summary

            # 3) If neither text nor images were flagged relevant, return text result or a fallback
            if combined_text_result:
                return False, combined_text_result
            else:
                return False, "No relevant content found in Word document."

        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    @staticmethod
    def _process_excel_document(
        excel_path: str,
        summary: str,
        description: str
    ) -> (bool, str):
        """
        Validate the Excel against the pre-agreed template:
        - mandatory table headers
        - applicant fields
        - dropdown 'Indicate Option'
        - SAP vs Table rows
        Returns (valid: bool, result_json: str).
        """
        APPLICANT_FIELDS = [
            "Name and Surname",
            "Date",
            "Requirement Description"
        ]
        MANDATORY_FILL_RGB = "FF002060"
        ALLOWED_OPTIONS = {
            "CREATION",
            "MODIFICATION/ACTUAL VALUE",
            "DELETION",
            "CHANGE",
            "ADDITION"
        }

        def load_grid(ws):
            """2D list of trimmed strings from each cell in the sheet."""
            return [
                [str(cell.value).strip() if cell.value is not None else "" for cell in row]
                for row in ws.iter_rows()
            ]

        def extract_component(grid):
            """Find first “component: XXX (YYY-ZZZ)” or the next cell if split."""
            for row in grid:
                for i, cell in enumerate(row):
                    if cell and "component" in cell.lower():
                        # after colon
                        if ":" in cell:
                            candidate = cell.split(":", 1)[1].strip()
                            if validate_component_format(candidate):
                                return candidate
                        for v in row[i + 1:]:
                            if v and validate_component_format(v.strip()):
                                return v.strip()
            return ""

        def validate_component_format(val):
            """Check format like “UDDC-65664” in parentheses and a dash."""
            if "(" in val and ")" in val:
                left, right = val.split("(", 1)
                code = right.rstrip(")").strip()
                return "-" in left and "-" in code
            return False

        def find_applicant_row(grid):
            for idx, row in enumerate(grid):
                if set(APPLICANT_FIELDS).issubset(set(row)):
                    return idx, {h: row.index(h) for h in APPLICANT_FIELDS}
            raise ValueError(f"Applicant headers not found: {APPLICANT_FIELDS}")

        def find_indicate(grid):
            for idx, row in enumerate(grid):
                for cell in row:
                    if cell in ALLOWED_OPTIONS:
                        return idx, cell
            return None, ""

        def find_value_row(grid, label):
            tgt = label.strip().lower()
            for idx, row in enumerate(grid):
                for cell in row:
                    if cell and cell.strip().lower() == tgt:
                        return idx
            raise ValueError(f"Row labeled '{label}' not found")

        def get_blue_header(ws):
            for r, row in enumerate(ws.iter_rows()):
                hdr = {}
                for c, cell in enumerate(row):
                    f = cell.fill
                    if f.fill_type == "solid":
                        rgb = getattr(f.start_color, "rgb", None) or getattr(f.fgColor, "rgb", None)
                        if rgb == MANDATORY_FILL_RGB and cell.value:
                            hdr[str(cell.value).strip()] = c
                if len(hdr) > 1:
                    return r, hdr
            raise ValueError("Mandatory-header row not found")

        def get_blue_rows(ws):
            blues = set()
            for r, row in enumerate(ws.iter_rows()):
                for cell in row:
                    f = cell.fill
                    if f.fill_type == "solid":
                        rgb = getattr(f.start_color, "rgb", None) or getattr(f.fgColor, "rgb", None)
                        if rgb == MANDATORY_FILL_RGB:
                            blues.add(r)
                            break
            return sorted(blues)

        result = {
            "status": "success",
            "component": "",
            "applicant": {},
            "indicate_option": "",
            "indicate_row": None,
            "header_row": None,
            "fields": {},
            "indicate_fields": {},
            "errors": []
        }

        try:
            wb = load_workbook(excel_path, data_only=True)
            ws = wb.active
            grid = load_grid(ws)

            # 0) Component
            comp = extract_component(grid)
            if comp:
                result["component"] = comp
            else:
                result["errors"].append("Missing Component field")
                result["status"] = "error"

            # 1) Applicant block
            try:
                a_idx, a_map = find_applicant_row(grid)
                for label, col in a_map.items():
                    v = grid[a_idx + 1][col] if a_idx + 1 < len(grid) else ""
                    result["applicant"][label] = v
                    if not v:
                        result["errors"].append(f"Applicant missing: {label}")
                        result["status"] = "error"
            except Exception as e:
                result["errors"].append(str(e))
                result["status"] = "error"

            # 2) Indicate Option
            i_idx, i_val = find_indicate(grid)
            if i_val:
                result["indicate_row"] = i_idx + 1
                result["indicate_option"] = i_val
            else:
                result["errors"].append("Missing Indicate Option")
                result["status"] = "error"

            # 3) Mandatory header
            try:
                hdr_idx, mand_map = get_blue_header(ws)
                result["header_row"] = hdr_idx + 1
            except Exception as e:
                result["errors"].append(str(e))
                result["status"] = "error"
                print(json.dumps(result, indent=2))

            # 4) Value rows (“CURRENT VALUE” & “NEW VALUE”), fallback to next two blue rows
            try:
                curr_idx = find_value_row(grid, "CURRENT VALUE")
                new_idx = find_value_row(grid, "NEW VALUE")
            except ValueError:
                blues = get_blue_rows(ws)
                after = [r for r in blues if r > hdr_idx]
                if len(after) >= 2:
                    curr_idx, new_idx = after[0], after[1]
                else:
                    result["errors"].append("Cannot locate CURRENT/NEW value rows")
                    result["status"] = "error"
                    print(json.dumps(result, indent=2))

            # 5) Extract each field’s [current, new]
            for fld, col in mand_map.items():
                cv = ws.cell(row=curr_idx + 1, column=col + 1).value or ""
                nv = ws.cell(row=new_idx + 1, column=col + 1).value or ""
                result["fields"][fld] = [str(cv).strip(), str(nv).strip()]

            # 7) Indicate-row fields
            if i_idx is not None:
                missing_ind = []
                ind_fields = {}
                for name, col in mand_map.items():
                    val = grid[i_idx][col] if col < len(grid[i_idx]) else ""
                    ind_fields[name] = val
                    if not val:
                        missing_ind.append(name)
                result["indicate_fields"] = ind_fields
                if missing_ind:
                    result["status"] = "error"
                    result["errors"].append(
                        "Indicate‐row missing mandatory columns: " + ", ".join(missing_ind)
                    )

        except Exception as e:
            result["status"] = "error"
            result["errors"].append(str(e))

        result_json = json.dumps(result, indent=2)
        is_valid = (result.get("status") == "success")
        return is_valid, result_json

    @staticmethod
    def _try_direct_multimodal(
        file_path: str,
        summary: str,
        description: str
    ) -> (bool, str):
        """
        Prompt the LLM directly with summary & description plus the file (image).
        Parse "Final Assessment: RELEVANT" vs. "NOT RELEVANT".
        """
        b64 = AttachmentAnalyzer.encode_file_to_base64(file_path)
        if not b64:
            raise Exception("Error encoding file to base64")
        mime = AttachmentAnalyzer._get_file_mime_type(file_path)

        prompt = (
            "You are a technical support agent analyzing an attached file to determine its relevance "
            "to the reported issue.\n\n"
            f"**Issue Summary:** {summary}\n\n"
            f"**Issue Description:** {description}\n\n"
            "**Your Task:**\n"
            "1. Analyze the attached content (image or document) in detail\n"
            "2. Compare the content with the issue summary and description\n"
            "3. Determine if the attachment is relevant to resolving this specific issue\n\n"
            "**Response Format:**\n"
            "1. **Description:** [Brief description of what you see]\n"
            "2. **Final Assessment:** **RELEVANT** or **NOT RELEVANT**\n"
            "3. **Support Reasoning:** [Why this decision helps in issue resolution]\n\n"
            "End your response with: TERMINATE"
        )

        llm, _ = LLM_Models.initialize_models()
        msg = HumanMessage(content=[
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
        ])
        res = llm.invoke([msg])
        ans = getattr(res, 'content', str(res)).strip()

        upper = ans.upper()
        if "FINAL ASSESSMENT: RELEVANT" in upper:
            return True, ans
        elif "FINAL ASSESSMENT: NOT RELEVANT" in upper:
            return False, ans
        else:
            return False, ans
